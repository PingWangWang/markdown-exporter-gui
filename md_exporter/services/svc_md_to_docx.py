#!/usr/bin/env python3
"""
Markdown -> DOCX 转换服务

本模块职责
----------
1) 将 Markdown 文本转换为 DOCX（基于 pandoc）。
2) 在未提供自定义模板时，执行统一的三步样式管线：
    - 删除不需要的段落/字符样式
    - 创建并刷新本模块定义的标准样式
    - 逐段应用字体、缩进、表格、图片、代码块等格式
3) 支持 Mermaid 代码块转图片，并在导出 DOCX 前替换为图片引用。

设计说明
--------
- 入口函数：`convert_md_to_docx()`
- 默认模板定位：`get_default_template()`
- 样式处理主流程：`_apply_formatting()` -> `_step1_*` / `_step2_*` / `_step3_*`
- 本模块避免直接暴露 docx XML 细节给上层调用者，相关复杂逻辑均封装在私有函数中。
"""

import re
import traceback
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor

from ..utils import get_logger
from ..utils.markdown_utils import get_md_text
from ..utils.mermaid_utils import replace_mermaid_with_images, cleanup_temp_images, extract_mermaid_blocks
from ..utils.pandoc_utils import pandoc_convert_file

logger = get_logger(__name__)

# ---------------------------------------------------------------------------
# 样式配置表
# - 每个字典对应一个“段落样式”定义
# - `style_keywords` 用于从 Pandoc 生成的样式名映射到本地配置
# - 带 `is_*` 标记的样式用于特殊分支（表格、图片、列表、代码块）
# ---------------------------------------------------------------------------
STYLE_CONFIGS: list[dict] = [
    {
        "name": "Heading 1",
        "style_keywords": ["Heading 1"],
        "font_name": "宋体",
        "font_name_latin": "Times New Roman",
        "font_size": Pt(20),
        "bold": True,
        "color": RGBColor(0, 0, 0),
        "first_line_indent": Pt(0),
        "left_indent": Pt(0),
        "line_spacing": 1.3,
        "space_before": Pt(3),
        "space_after": Pt(3),
    },
    {
        "name": "Heading 2",
        "style_keywords": ["Heading 2"],
        "font_name": "宋体",
        "font_name_latin": "Times New Roman",
        "font_size": Pt(18),
        "bold": True,
        "color": RGBColor(0, 0, 0),
        "first_line_indent": Pt(0),
        "left_indent": Pt(0),
        "line_spacing": 1.3,
        "space_before": Pt(3),
        "space_after": Pt(3),
    },
    {
        "name": "Heading 3",
        "style_keywords": ["Heading 3"],
        "font_name": "宋体",
        "font_name_latin": "Times New Roman",
        "font_size": Pt(16),
        "bold": True,
        "color": RGBColor(0, 0, 0),
        "first_line_indent": Pt(0),
        "left_indent": Pt(0),
        "line_spacing": 1.3,
        "space_before": Pt(3),
        "space_after": Pt(3),
    },
    {
        "name": "Heading 4",
        "style_keywords": ["Heading 4"],
        "font_name": "宋体",
        "font_name_latin": "Times New Roman",
        "font_size": Pt(14),
        "bold": True,
        "color": RGBColor(0, 0, 0),
        "first_line_indent": Pt(0),
        "left_indent": Pt(0),
        "line_spacing": 1.3,
        "space_before": Pt(3),
        "space_after": Pt(3),
    },
    {
        "name": "Heading 5",
        "style_keywords": ["Heading 5"],
        "font_name": "宋体",
        "font_name_latin": "Times New Roman",
        "font_size": Pt(12),
        "bold": True,
        "color": RGBColor(0, 0, 0),
        "first_line_indent": Pt(0),
        "left_indent": Pt(0),
        "line_spacing": 1.3,
        "space_before": Pt(3),
        "space_after": Pt(3),
    },
    {
        "name": "Heading 6",
        "style_keywords": ["Heading 6"],
        "font_name": "宋体",
        "font_name_latin": "Times New Roman",
        "font_size": Pt(12),
        "bold": True,
        "color": RGBColor(0, 0, 0),
        "first_line_indent": Pt(0),
        "left_indent": Pt(0),
        "line_spacing": 1.3,
        "space_before": Pt(3),
        "space_after": Pt(3),
    },
    {
        "name": "Normal",
        "style_keywords": ["Normal"],
        "font_name": "宋体",
        "font_name_latin": "Times New Roman",
        "font_size": Pt(12),
        "bold": False,
        "color": RGBColor(0, 0, 0),
        "first_line_indent": Pt(24),
        "left_indent": Pt(0),
        "line_spacing": 1.3,
        "space_before": Pt(0),
        "space_after": Pt(0),
    },
    {
        "name": "List Paragraph",
        "style_keywords": ["List Paragraph", "List"],
        "font_name": "宋体",
        "font_name_latin": "Times New Roman",
        "font_size": Pt(12),
        "bold": False,
        "color": RGBColor(0, 0, 0),
        "first_line_indent": Pt(0),
        "left_indent": Pt(0),
        "line_spacing": 1.3,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "is_list": True,
    },
    {
        "name": "Custom List",
        "style_keywords": [],
        "font_name": "宋体",
        "font_name_latin": "Times New Roman",
        "font_size": Pt(12),
        "bold": False,
        "color": RGBColor(0, 0, 0),
        "first_line_indent": Pt(24),
        "left_indent": Pt(0),
        "line_spacing": 1.3,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "is_custom_list": True,
    },
    {
        "name": "Table Text",
        "style_keywords": [],
        "font_name": "宋体",
        "font_name_latin": "Times New Roman",
        "font_size": Pt(12),
        "bold": False,
        "color": RGBColor(0, 0, 0),
        "first_line_indent": Pt(0),
        "left_indent": Pt(0),
        "line_spacing": 1.0,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "alignment": 1,
        "is_table": True,
    },
    {
        "name": "Image Paragraph",
        "style_keywords": [],
        "font_name": "宋体",
        "font_name_latin": "Times New Roman",
        "font_size": Pt(12),
        "bold": False,
        "color": RGBColor(0, 0, 0),
        "first_line_indent": Pt(0),
        "left_indent": Pt(0),
        "line_spacing": 1.3,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "alignment": 1,
        "is_image": True,
    },
    {
        "name": "Code Block",
        "style_keywords": [],
        "font_name": "宋体",
        "font_name_latin": "Times New Roman",
        "font_size": Pt(10.5),  # 五号字
        "bold": False,
        "italic": True,  # 斜体
        "color": RGBColor(255, 255, 255),  # 更明显的白色
        "first_line_indent": Pt(0),
        "left_indent": Pt(24),  # 左缩进2字符（所有行都缩进）
        "line_spacing": 1.0,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "background_color": RGBColor(0, 0, 0),  # 黑色背景
        "is_code": True,
    },
]

REQUIRED_PARAGRAPH_STYLES: set[str] = {
    "Normal",
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "Heading 4",
    "Heading 5",
    "Heading 6",
    "Table Text",
    "List Paragraph",
    "Image Paragraph",
    "Custom List",
    "Code Block",
    "Source Code",  # Pandoc generated code block style
    "Preformatted Text",  # Pandoc generated code block style
}

REQUIRED_CHARACTER_STYLES: set[str] = {
    "Default Paragraph Font",
    "Hyperlink",
    "Strong",
    "Emphasis",
}

# 常用样式配置快捷引用，避免运行时重复查找
_NORMAL_CONFIG: dict = next(c for c in STYLE_CONFIGS if c["name"] == "Normal")
_TABLE_CONFIG: dict = next(c for c in STYLE_CONFIGS if c.get("is_table"))
_IMAGE_CONFIG: dict = next(c for c in STYLE_CONFIGS if c.get("is_image"))
_CUSTOM_LIST_CONFIG: dict = next(c for c in STYLE_CONFIGS if c.get("is_custom_list"))
_CODE_CONFIG: dict = next(c for c in STYLE_CONFIGS if c.get("is_code"))

# Pandoc/语法高亮可能产生的代码样式关键字
CODE_STYLE_KEYWORDS: tuple[str, ...] = (
    "Preformatted",
    "Code",
    "Source Code",
    "NormalTok",
    "Verbatim",
    "KeywordTok",
    "StringTok",
    "CommentTok",
    "FunctionTok",
    "VariableTok",
    "DataTypeTok",
    "DecValTok",
    "BaseNTok",
    "FloatTok",
    "ConstantTok",
    "CharTok",
    "SpecialCharTok",
    "ImportTok",
    "DocumentationTok",
    "AnnotationTok",
    "OtherTok",
    "ControlFlowTok",
    "OperatorTok",
    "BuiltInTok",
    "ExtensionTok",
    "PreprocessorTok",
    "AttributeTok",
    "RegionMarkerTok",
    "InformationTok",
    "WarningTok",
    "AlertTok",
    "ErrorTok",
)

# 常量区：图像尺寸换算与页面默认值
# 1pt = 12700 EMU（Office Open XML 单位）
EMU_PER_POINT = 12700.0
DEFAULT_PAGE_WIDTH_PT = 595
DEFAULT_PAGE_HEIGHT_PT = 842
IMAGE_SIDE_MARGIN_PT = 36
IMAGE_TOP_BOTTOM_MARGIN_PT = 72


# ---------------------------------------------------------------------------
# 模块级辅助函数
# ---------------------------------------------------------------------------


def _get_config_for_style(style_name: str) -> dict:
    """根据段落样式名获取配置。

    说明：
    - 按 `STYLE_CONFIGS` 顺序匹配 `style_keywords`。
    - 表格样式由专门分支处理，这里跳过 `is_table`。
    - 未命中时回退到 `_NORMAL_CONFIG`。
    """
    for config in STYLE_CONFIGS:
        if config.get("is_table"):
            continue
        for keyword in config["style_keywords"]:
            if keyword in style_name:
                return config
    return _NORMAL_CONFIG


def _set_rfonts(rpr_elem, font_name: str, font_name_latin: str | None = None) -> None:
    """设置 run 的中西文字体，并移除主题字体覆盖。

    Word 中字体可能由主题（theme）覆盖显式配置；为保证导出一致性，
    这里主动移除 `asciiTheme/hAnsiTheme/themeEastAsia/cstheme`。
    """
    rFonts = rpr_elem.get_or_add_rFonts()
    # Set Latin fonts (ascii and hAnsi)
    latin_font = font_name_latin if font_name_latin else font_name
    rFonts.set(qn("w:ascii"), latin_font)
    rFonts.set(qn("w:hAnsi"), latin_font)
    # Set East Asian fonts
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:themeEastAsia", "w:cstheme"):
        rFonts.attrib.pop(qn(attr), None)


# Matches paragraphs that start with a number/letter list marker and should not be indented
_NO_INDENT_PATTERN = re.compile(r"^\s*(\d+[.、）)）]|[（(]\d+[）)]|[一二三四五六七八九十百]+[、.]|[a-zA-Z][.)])")


def _needs_no_indent(paragraph) -> bool:
    """判断段落是否应取消首行缩进。

    当前策略：仅当段落文本以编号/字母序号开头时取消缩进，
    其他普通段落仍保留 `Normal` 的首行缩进规则。
    """
    return bool(_NO_INDENT_PATTERN.match(paragraph.text))


def _has_num_pr(paragraph) -> bool:
    """判断段落是否包含 `w:numPr`（通常表示列表项）。"""
    pPr = paragraph._p.find(qn("w:pPr"))
    return pPr is not None and pPr.find(qn("w:numPr")) is not None


def _has_image(paragraph) -> bool:
    """判断段落中是否包含图片节点。"""
    # 通过 drawing/pic 节点快速判断（兼容内联/浮动图）
    for child in paragraph._element.iter():
        if child.tag.endswith("drawing") or child.tag.endswith("pic"):
            return True
    return False


def _is_code_block(paragraph) -> bool:
    """判断段落是否为代码块。

    判定顺序：
    1) 样式名是否包含代码关键词（Pandoc 常见输出）；
    2) run 字体是否为等宽字体（Consolas/Courier/Monospace）。
    """
    style_name = paragraph.style.name if paragraph.style else ""
    for keyword in CODE_STYLE_KEYWORDS:
        if keyword in style_name:
            logger.debug(f"Detected code block by style name: {style_name}")
            return True
    
    # 兜底：检测等宽字体
    for run in paragraph.runs:
        font_name = run.font.name
        if font_name and ("Consolas" in font_name or "Courier" in font_name or "Monospace" in font_name):
            logger.debug(f"Detected code block by font: {font_name}")
            return True
    
    return False


def _get_image_limits(doc) -> tuple[float, float]:
    """根据文档页尺寸计算图片可用宽高上限（pt）。"""
    section = doc.sections[0] if doc.sections else None
    page_width = section.page_width.pt if section else DEFAULT_PAGE_WIDTH_PT
    page_height = section.page_height.pt if section else DEFAULT_PAGE_HEIGHT_PT
    max_width = page_width - 2 * IMAGE_SIDE_MARGIN_PT
    max_height = page_height - 2 * IMAGE_TOP_BOTTOM_MARGIN_PT
    return max_width, max_height


def _xml_tag_name(elem) -> str:
    """提取 XML 本地标签名（去掉命名空间前缀）。"""
    return elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag


def _get_extent_size_emu(extent_elem) -> tuple[int, int] | None:
    """读取 `wp:extent` 的宽高（EMU）。读取失败返回 `None`。"""
    cx_val = extent_elem.get(qn("wp:cx")) or extent_elem.get("cx")
    cy_val = extent_elem.get(qn("wp:cy")) or extent_elem.get("cy")
    if not cx_val or not cy_val:
        return None
    return int(cx_val), int(cy_val)


def _set_extent_size_emu(extent_elem, width_emu: int, height_emu: int) -> None:
    """写入 `wp:extent` 的宽高（EMU），兼容命名空间/非命名空间属性。"""
    if extent_elem.get(qn("wp:cx")) is not None:
        extent_elem.set(qn("wp:cx"), str(width_emu))
        extent_elem.set(qn("wp:cy"), str(height_emu))
    else:
        extent_elem.set("cx", str(width_emu))
        extent_elem.set("cy", str(height_emu))


def _scale_extent_if_needed(extent_elem, context_tag: str, max_width_pt: float, max_height_pt: float) -> None:
    """按比例缩放单个图片节点，使其不超过页面可用范围。"""
    size_emu = _get_extent_size_emu(extent_elem)
    if size_emu is None:
        logger.warning("DEBUG: cx_val or cy_val is None/empty")
        return

    current_width_emu, current_height_emu = size_emu
    current_width_pt = current_width_emu / EMU_PER_POINT
    current_height_pt = current_height_emu / EMU_PER_POINT

    logger.info(
        "Image (%s) size: %.0fx%.0fpt, max: %.0fx%.0fpt",
        context_tag,
        current_width_pt,
        current_height_pt,
        max_width_pt,
        max_height_pt,
    )

    scale_factor = 1.0
    if current_width_pt > max_width_pt:
        scale_factor = min(scale_factor, max_width_pt / current_width_pt)
    if current_height_pt > max_height_pt:
        scale_factor = min(scale_factor, max_height_pt / current_height_pt)

    if scale_factor >= 1.0:
        logger.info("Image within limits, no scaling needed")
        return

    new_width_emu = int(current_width_emu * scale_factor)
    new_height_emu = int(current_height_emu * scale_factor)
    _set_extent_size_emu(extent_elem, new_width_emu, new_height_emu)

    logger.info(
        "✓ Scaled image from %.0fx%.0fpt to %.0fx%.0fpt",
        current_width_pt,
        current_height_pt,
        new_width_emu / EMU_PER_POINT,
        new_height_emu / EMU_PER_POINT,
    )


def _scale_images_in_paragraph(paragraph, max_width_pt: float, max_height_pt: float) -> None:
    """扫描并缩放段落内全部图片节点。"""
    logger.info("Found image paragraph, checking scaling...")
    current_context_tag = "unknown"

    try:
        for elem in paragraph._element.iter():
            tag_name = _xml_tag_name(elem)

            if tag_name in {"inline", "anchor"}:
                current_context_tag = tag_name
                continue

            if tag_name != "extent":
                continue

            try:
                _scale_extent_if_needed(elem, current_context_tag, max_width_pt, max_height_pt)
            except Exception as img_exc:
                logger.warning(f"Failed to scale image: {img_exc}")
                logger.warning(f"Traceback: {traceback.format_exc()}")
    except Exception as exc:
        logger.error(f"Error iterating paragraph elements: {exc}")
        logger.error(f"Traceback: {traceback.format_exc()}")


def _apply_table_layout(tbl) -> None:
    """为表格统一设置宽度与边框样式。

    - 宽度：100% 页面宽（`w:type=pct, w=5000`）
    - 边框：外边框与内部网格均为 single
    """
    tblPr = tbl._element.tblPr
    if tblPr is None:
        tblPr = parse_xml(f"<w:tblPr {nsdecls('w')}/>")
        tbl._element.insert(0, tblPr)

    existing_tblW = tblPr.find(qn("w:tblW"))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)

    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
    tblPr.append(tblW)

    tblBorders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
    </w:tblBorders>'''
    tblBorders = parse_xml(tblBorders_xml)

    existing_tblBorders = tblPr.find(qn("w:tblBorders"))
    if existing_tblBorders is not None:
        tblPr.remove(existing_tblBorders)

    tblPr.append(tblBorders)


def _set_cell_vertical_center(cell) -> None:
    """设置单元格垂直居中。"""
    tc = cell._element
    tcPr = tc.tcPr if tc.tcPr is not None else parse_xml(f'<w:tcPr {nsdecls("w")}/>')
    if tc.tcPr is None:
        tc.insert(0, tcPr)

    existing_vAlign = tcPr.find(qn("w:vAlign"))
    if existing_vAlign is not None:
        tcPr.remove(existing_vAlign)

    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)


def _format_table_content(doc) -> None:
    """对文档中所有表格应用统一格式。

    包含：
    - 表格级（宽度、边框）
    - 单元格级（垂直居中）
    - 段落级（`Table Text` 样式 + run 字体设置）
    """
    table_text_style = doc.styles["Table Text"]
    for tbl in doc.tables:
        _apply_table_layout(tbl)
        for row in tbl.rows:
            for cell in row.cells:
                _set_cell_vertical_center(cell)
                for para in cell.paragraphs:
                    try:
                        para.style = table_text_style
                        _apply_para_formatting(para, _TABLE_CONFIG, is_table=True)
                    except Exception as exc:
                        logger.warning(f"Failed to format table cell: {exc}")


def _apply_para_formatting(paragraph, config: dict, is_table: bool = False) -> None:
    """按配置应用段落与 run 级格式。

    参数：
    - `paragraph`: 待处理段落
    - `config`: 来自 `STYLE_CONFIGS` 的样式配置
    - `is_table`: 是否处于表格上下文（用于对齐与缩进策略）
    """
    pf = paragraph.paragraph_format
    pf.line_spacing = config["line_spacing"]
    pf.space_before = config["space_before"]
    pf.space_after = config["space_after"]
    
    # 代码块：尽量保持段内连续，避免分页断开阅读
    if config.get("is_code"):
        pf.keep_together = True
        pf.keep_with_next = True
    
    # 自定义列表：显式覆盖 Pandoc 生成的缩进
    if config.get("is_custom_list"):
        pf.left_indent = config["left_indent"]
        pf.first_line_indent = config["first_line_indent"]
    # 列表项：编号系统负责部分缩进，这里仅控制左缩进基线
    elif config.get("is_list"):
        # 首行缩进固定为 0，避免与编号悬挂缩进冲突
        pf.left_indent = config["left_indent"]
        pf.first_line_indent = Pt(0)
    elif not _has_num_pr(paragraph):
        if config["first_line_indent"] and not is_table and _needs_no_indent(paragraph):
            pf.first_line_indent = Pt(0)
        else:
            pf.first_line_indent = config["first_line_indent"]
            pf.left_indent = config["left_indent"]
    if is_table and "alignment" in config:
        pf.alignment = config["alignment"]
    # 图片段落：保险起见再次设置为居中
    if config.get("is_image"):
        pf.alignment = 1  # Center alignment

    for run in paragraph.runs:
        run.font.color.rgb = config["color"]
        run.font.name = config.get("font_name_latin", config["font_name"])
        run.font.size = config["font_size"]
        # 保留 Pandoc 显式粗体（如 Markdown 的 **text**）
        if config["bold"]:
            run.font.bold = True
        elif run.font.bold is not True:
            run.font.bold = config["bold"]
        
        # 配置要求斜体时强制设置
        if config.get("italic"):
            run.font.italic = True
        
        # 代码块统一使用西文字体（避免等宽字体被中文字体覆盖）
        if config.get("is_code"):
            _set_rfonts(
                run._element.get_or_add_rPr(),
                config.get("font_name_latin", config["font_name"]),
            )
        else:
            _set_rfonts(
                run._element.get_or_add_rPr(),
                config["font_name"],
                config.get("font_name_latin"),
            )


# ---------------------------------------------------------------------------
# 三步格式化管线
# ---------------------------------------------------------------------------


def _step1_delete_styles(doc) -> None:
    """步骤 1：删除不在白名单中的样式，并做必要迁移。

    关键点：
    - 段落样式：不在 `REQUIRED_PARAGRAPH_STYLES` 中的将被移除；
    - 字符样式：不在 `REQUIRED_CHARACTER_STYLES` 中的将被移除；
    - 删除前先把引用它们的段落/run 重映射，避免悬空引用。
    """
    normal_style = doc.styles["Normal"]
    existing_names = {s.name for s in doc.styles}
    # List Paragraph may not exist yet before _step2 creates it; create a minimal
    # placeholder now so that list sub-styles can be reassigned to it.
    if "List Paragraph" not in existing_names:
        list_style = doc.styles.add_style("List Paragraph", 1)
        list_style.base_style = normal_style
    else:
        list_style = doc.styles["List Paragraph"]

    styles_to_delete = [
        s
        for s in doc.styles
        if (s.type == 1 and s.name not in REQUIRED_PARAGRAPH_STYLES)
        or (s.type == 2 and s.name not in REQUIRED_CHARACTER_STYLES)
    ]

    for style_obj in styles_to_delete:
        try:
            if style_obj.type == 1:
                # Pandoc generates "List Paragraph 1/2/..." for nested lists;
                # reassign those to List Paragraph so they keep bullet formatting.
                is_list_sub = "List" in style_obj.name
                fallback_style = list_style if is_list_sub else normal_style
                for para in doc.paragraphs:
                    if para.style.name == style_obj.name:
                        para.style = fallback_style
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if para.style.name == style_obj.name:
                                    para.style = fallback_style
            else:
                style_id = style_obj._element.get(qn("w:styleId"))
                all_paras = list(doc.paragraphs)
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            all_paras.extend(cell.paragraphs)
                for para in all_paras:
                    for run in para.runs:
                        rPr = run._element.find(qn("w:rPr"))
                        if rPr is not None:
                            rStyle = rPr.find(qn("w:rStyle"))
                            if rStyle is not None and rStyle.get(qn("w:val")) == style_id:
                                rPr.remove(rStyle)

            style_obj._element.getparent().remove(style_obj._element)
            logger.info(f"Deleted style: {style_obj.name}")
        except Exception as exc:
            logger.warning(f"Failed to delete style '{style_obj.name}': {exc}")


def _step2_create_styles(doc) -> None:
    """步骤 2：创建/刷新本模块定义的标准样式。"""
    for config in STYLE_CONFIGS:
        name = config["name"]
        try:
            existing_names = {s.name for s in doc.styles}
            style = doc.styles[name] if name in existing_names else doc.styles.add_style(name, 1)

            style.font.name = config.get("font_name_latin", config["font_name"])
            style.font.size = config["font_size"]
            style.font.bold = config["bold"]
            style.font.color.rgb = config["color"]

            # 使用 style.element.get_or_add_rPr() 获取 CT_RPr 更稳定；
            # style.font._element 在部分场景会返回 CT_Style。
            rPr = style.element.get_or_add_rPr()
            _set_rfonts(rPr, config["font_name"], config.get("font_name_latin"))

            # 移除主题色覆盖，确保颜色按显式配置生效
            color_elem = rPr.find(qn("w:color"))
            if color_elem is not None:
                for attr in ("w:themeColor", "w:themeShade", "w:themeTint"):
                    color_elem.attrib.pop(qn(attr), None)

            pf = style.paragraph_format
            pf.line_spacing = config["line_spacing"]
            pf.first_line_indent = config["first_line_indent"]
            if config["left_indent"] is not None:
                pf.left_indent = config["left_indent"]
            pf.space_before = config["space_before"]
            pf.space_after = config["space_after"]
            
            # 代码块背景色（段落底纹）
            if config.get("background_color"):
                pPr = style.element.get_or_add_pPr()
                shd = pPr.find(qn("w:shd"))
                if shd is None:
                    shd_xml = '<w:shd {} w:val="clear"/>'.format(nsdecls("w"))
                    shd = parse_xml(shd_xml)
                    pPr.append(shd)
                bg_color = config["background_color"]
                hex_color = f"{bg_color[0]:02X}{bg_color[1]:02X}{bg_color[2]:02X}"
                shd.set(qn("w:fill"), hex_color)
        except Exception as exc:
            logger.warning(f"Failed to create/update style '{name}': {exc}")


def _step3_apply_to_content(doc) -> None:
    """步骤 3：遍历正文并按内容类型应用样式。"""
    image_para_style = doc.styles["Image Paragraph"]
    custom_list_style = doc.styles["Custom List"]
    code_block_style = doc.styles["Code Block"]

    code_block_count = 0
    max_image_width, max_image_height = _get_image_limits(doc)

    for para in doc.paragraphs:
        try:
            if _is_code_block(para):
                para.style = code_block_style
                _apply_para_formatting(para, _CODE_CONFIG)
                code_block_count += 1
                logger.debug(f"Applied Code Block style to paragraph: {para.text[:50]}")
            elif _has_image(para):
                para.style = image_para_style
                _apply_para_formatting(para, _IMAGE_CONFIG)
                _scale_images_in_paragraph(para, max_image_width, max_image_height)
            elif _has_num_pr(para):
                para.style = custom_list_style
                _apply_para_formatting(para, _CUSTOM_LIST_CONFIG)
            else:
                style_name = para.style.name if para.style else "Normal"
                _apply_para_formatting(para, _get_config_for_style(style_name))
        except Exception as exc:
            logger.warning(f"Failed to format paragraph: {exc}")

    logger.info(f"Total code blocks formatted: {code_block_count}")
    _format_table_content(doc)


def _build_pandoc_extra_args(
    template_path: Path | None,
    is_enable_toc: bool,
    resource_paths: list[str] | None = None,
) -> list[str]:
    """构建 Pandoc 参数列表。

    包含：
    - `--reference-doc`：参考模板
    - `--toc`：目录开关
    - `--resource-path`：资源搜索路径（图片等）
    """
    extra_args: list[str] = []
    if template_path and template_path.exists():
        extra_args.append(f"--reference-doc={template_path}")
    if is_enable_toc:
        extra_args.append("--toc")
    if resource_paths:
        extra_args.append(f"--resource-path={';'.join(resource_paths)}")
    return extra_args


def _convert_markdown_file_to_docx(
    source_md_file: Path,
    output_path: Path,
    template_path: Path | None,
    is_enable_toc: bool,
    resource_paths: list[str] | None = None,
) -> None:
    """将单个 Markdown 文件转换为 DOCX，并按需应用默认格式化。"""
    final_template_path = template_path or get_default_template()
    extra_args = _build_pandoc_extra_args(final_template_path, is_enable_toc, resource_paths)

    pandoc_convert_file(
        source_file=str(source_md_file),
        input_format="markdown",
        dest_format="docx",
        outputfile=str(output_path),
        extra_args=extra_args,
    )

    if not template_path:
        _apply_formatting(output_path)
    else:
        logger.info(f"Using custom template, skipping formatting: {template_path}")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def _apply_formatting(docx_path: Path) -> None:
    """
        对输出 DOCX 应用三步格式化管线：
        1) 删除非白名单样式；
        2) 依据 `STYLE_CONFIGS` 刷新标准样式；
        3) 遍历正文/表格并应用段落与 run 级格式。
    """
    doc = Document(docx_path)
    _step1_delete_styles(doc)
    _step2_create_styles(doc)
    _step3_apply_to_content(doc)
    doc.save(docx_path)
    logger.info(f"Applied formatting to {docx_path}")


def convert_md_to_docx(
    md_text: str,
    output_path: Path,
    template_path: Path | None = None,
    is_strip_wrapper: bool = False,
    is_enable_toc: bool = False,
    convert_mermaid: bool = True,
    save_mermaid_images: bool = False,
    output_dir: Path | None = None,
) -> None:
    """
    将 Markdown 文本转换为 DOCX。

    Args:
        md_text: 输入 Markdown 文本
        output_path: 输出 DOCX 路径
        template_path: 自定义 DOCX 模板（可选）
        is_strip_wrapper: 是否剥离代码块包装
        is_enable_toc: 是否启用目录
        convert_mermaid: 是否启用 Mermaid 转图片
        save_mermaid_images: 是否将 Mermaid 图片持久保存到输出目录
        output_dir: Mermaid 图片输出目录（当 `save_mermaid_images=True` 时建议提供）

    Raises:
        ValueError: 输入处理失败
        Exception: 转换或后处理失败
    """
    processed_md = get_md_text(md_text, is_strip_wrapper=is_strip_wrapper)

    # 预扫描 Mermaid 代码块，用于分流到“图表转换流程”或“标准流程”
    mermaid_blocks = extract_mermaid_blocks(processed_md)

    if mermaid_blocks and convert_mermaid:
        logger.info(f"检测到 {len(mermaid_blocks)} 个 Mermaid 图表，开始转换...")

        # 根据保存策略决定图片落盘位置
        if save_mermaid_images and output_dir:
            # 创建图片保存目录
            images_dir = output_dir / "mermaid_images"
            images_dir.mkdir(exist_ok=True)
            save_path = images_dir
            logger.info(f"Mermaid 图片将保存到: {images_dir}")
        else:
            # 不保存到目标目录：统一写临时目录，流程结束后清理
            save_path = None

        # 统一使用临时目录承载中间产物（Markdown、转换图片等）
        with TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)

            # 若不持久化 Mermaid 图片，则直接写入临时目录
            image_save_path = save_path if save_mermaid_images else temp_path

            # 替换 Mermaid 代码块为图片引用（使用 PNG 格式，通过 scale 参数提高清晰度）
            modified_md, generated_images, mermaid_stats = replace_mermaid_with_images(
                processed_md,
                image_save_path,
                image_format="png",
                timeout=15,  # 增加超时时间，因为大图片需要更长时间
                max_retries=3,
                retry_delay=2,
                scale=3  # 3倍缩放提高清晰度
            )

            # 输出 Mermaid 转换统计，便于排查失败图表
            if mermaid_stats["total"] > 0:
                logger.info("=" * 50)
                logger.info(f"Mermaid 转换汇总:")
                logger.info(f"  总计: {mermaid_stats['total']} 个")
                logger.info(f"  成功: {mermaid_stats['success']} 个")
                if mermaid_stats["failed"] > 0:
                    logger.info(f"  失败: {mermaid_stats['failed']} 个")
                logger.info("=" * 50)

            # 将替换后的 Markdown（Mermaid -> ![](...)）写到临时文件
            temp_md_file = temp_path / "temp.md"
            temp_md_file.write_text(modified_md, encoding="utf-8")

            # 指定 Pandoc 资源路径，确保图片可被解析
            resource_paths = [str(temp_path)]
            if save_mermaid_images and save_path:
                resource_paths.append(str(save_path))

            try:
                _convert_markdown_file_to_docx(
                    source_md_file=temp_md_file,
                    output_path=output_path,
                    template_path=template_path,
                    is_enable_toc=is_enable_toc,
                    resource_paths=resource_paths,
                )
            finally:
                # 非持久化模式：主动清理 Mermaid 图片临时文件
                if not save_mermaid_images:
                    cleanup_temp_images(generated_images)
                    logger.info("已清理临时图片文件")
                else:
                    logger.info(f"已保存 {len(generated_images)} 个 Mermaid 图片到: {save_path}")
    else:
        # 无 Mermaid：走标准 Markdown -> DOCX 流程
        logger.info("未检测到 Mermaid 图表，使用标准转换流程")

        with TemporaryDirectory() as temp_dir:
            temp_md_file = Path(temp_dir) / "temp.md"
            temp_md_file.write_text(processed_md, encoding="utf-8")
            _convert_markdown_file_to_docx(
                source_md_file=temp_md_file,
                output_path=output_path,
                template_path=template_path,
                is_enable_toc=is_enable_toc,
            )


def get_default_template() -> Path | None:
    """返回内置 DOCX 模板路径；若不存在则返回 `None`。"""
    template = Path(__file__).resolve().parent.parent / "assets" / "template" / "docx_template.docx"
    if template.exists():
        return template
    logger.warning(f"Default DOCX template not found at {template}")
    return None
